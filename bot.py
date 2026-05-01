import os
import logging
import asyncio
from datetime import datetime
from typing import Dict, Any

from aiogram import Bot, Dispatcher, types
from aiogram.contrib.middlewares.logging import LoggingMiddleware
from aiogram.dispatcher import FSMContext
from aiogram.dispatcher.filters.state import State, StatesGroup
from aiogram.types import ParseMode, InputFile, ReplyKeyboardMarkup, KeyboardButton, InlineKeyboardMarkup, InlineKeyboardButton
from aiogram.utils import executor
from aiogram.contrib.fsm_storage.sqlite import SQLiteStorage

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from aiohttp import web

# ----------------------------------------------------------------------
# НАСТРОЙКИ
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

BOT_TOKEN = os.getenv("BOT_TOKEN")
if not BOT_TOKEN:
    raise ValueError("Не задан BOT_TOKEN в переменных окружения")

storage = SQLiteStorage(database="fsm.db")
bot = Bot(token=BOT_TOKEN, parse_mode=ParseMode.HTML)
dp = Dispatcher(bot, storage=storage)
dp.middleware.setup(LoggingMiddleware())

# ----------------------------------------------------------------------
# СОСТОЯНИЯ (FSM)
class Form(StatesGroup):
    full_name = State()
    position = State()
    department = State()
    date = State()
    experience = State()
    supervisor = State()
    mission_goal = State()
    mission_impact = State()
    mission_absence = State()
    key_functions = State()
    competencies = State()
    direct_value = State()
    indirect_value = State()
    problems = State()
    improvements = State()
    goals = State()
    feedback = State()
    self_rating = State()
    self_strength = State()
    self_weakness = State()
    self_one_change = State()

# Хранилище ответов пользователей
user_data: Dict[int, Dict[str, Any]] = {}

# ----------------------------------------------------------------------
# КЛАВИАТУРЫ
def get_main_keyboard():
    keyboard = ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=False)
    keyboard.add(KeyboardButton("🏁 Старт"), KeyboardButton("📝 Заполнить анкету"))
    keyboard.add(KeyboardButton("❌ Отмена"), KeyboardButton("ℹ️ Помощь"))
    return keyboard

def get_cancel_keyboard():
    keyboard = ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=False)
    keyboard.add(KeyboardButton("❌ Отмена"))
    return keyboard

def get_rating_keyboard():
    keyboard = InlineKeyboardMarkup(row_width=5)
    buttons = [InlineKeyboardButton(str(i), callback_data=f"rating_{i}") for i in range(1, 11)]
    keyboard.add(*buttons)
    return keyboard

def get_date_keyboard():
    keyboard = InlineKeyboardMarkup(row_width=2)
    today = datetime.now().strftime("%d.%m.%Y")
    keyboard.add(
        InlineKeyboardButton(f"Сегодня ({today})", callback_data="date_today"),
        InlineKeyboardButton("Ввести вручную", callback_data="date_manual")
    )
    return keyboard

# ----------------------------------------------------------------------
# ОБРАБОТЧИКИ ОТМЕНЫ (должны быть ПЕРВЫМИ)
@dp.message_handler(commands=['cancel'], state='*')
async def cmd_cancel(message: types.Message, state: FSMContext):
    current_state = await state.get_state()
    if current_state is None:
        await message.answer("Нет активного опроса.", reply_markup=get_main_keyboard())
        return
    await state.finish()
    if message.from_user.id in user_data:
        del user_data[message.from_user.id]
    await message.answer("❌ Опрос отменён. Чтобы начать заново, нажмите «📝 Заполнить анкету».", reply_markup=get_main_keyboard())

@dp.message_handler(lambda message: message.text == "❌ Отмена", state='*')
async def cancel_button_handler(message: types.Message, state: FSMContext):
    await cmd_cancel(message, state)

# ----------------------------------------------------------------------
# ОСНОВНЫЕ КОМАНДЫ И КНОПКИ
@dp.message_handler(commands=['start'])
async def cmd_start(message: types.Message):
    await message.answer(
        "🏢 <b>Бот для заполнения анкеты сотрудника</b>\n\n"
        "Я задам вопросы о вашей работе и сформирую документ .docx.\n\n"
        "👉 Нажмите <b>«📝 Заполнить анкету»</b>, чтобы начать.\n"
        "👉 <b>«❌ Отмена»</b> – прервать опрос.\n"
        "👉 <b>«ℹ️ Помощь»</b> – справка.",
        reply_markup=get_main_keyboard()
    )

@dp.message_handler(commands=['help'])
async def cmd_help(message: types.Message):
    await message.answer(
        "📌 <b>Доступные действия</b>:\n\n"
        "🏁 Старт – приветственное сообщение\n"
        "📝 Заполнить анкету – начать новый опрос\n"
        "❌ Отмена – прервать текущее анкетирование\n"
        "ℹ️ Помощь – это сообщение\n\n"
        "Для оценки (1-10) будут появляться кнопки.",
        reply_markup=get_main_keyboard()
    )

@dp.message_handler(commands=['fill'])
async def cmd_fill(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    current_state = await state.get_state()
    if current_state is not None:
        await message.answer("У вас уже идёт опрос. Используйте «❌ Отмена», чтобы начать заново.", reply_markup=get_cancel_keyboard())
        return
    user_data[user_id] = {}
    await message.answer(
        "✏️ <b>Начинаем заполнение анкеты.</b>\nВведите ваше <b>ФИО</b>:",
        reply_markup=get_cancel_keyboard()
    )
    await Form.full_name.set()

@dp.message_handler(lambda message: message.text == "🏁 Старт")
async def russian_start(message: types.Message):
    await cmd_start(message)

@dp.message_handler(lambda message: message.text == "📝 Заполнить анкету")
async def russian_fill(message: types.Message, state: FSMContext):
    await cmd_fill(message, state)

@dp.message_handler(lambda message: message.text == "ℹ️ Помощь")
async def russian_help(message: types.Message):
    await cmd_help(message)

# ----------------------------------------------------------------------
# ОБРАБОТЧИК СООБЩЕНИЙ ВНЕ СОСТОЯНИЙ
@dp.message_handler(state=None)
async def no_state_handler(message: types.Message):
    await message.answer(
        "❗ Чтобы начать анкетирование, нажмите кнопку «📝 Заполнить анкету».\n"
        "Если вы уже начали, но бот не отвечает – возможно, сессия сбросилась.\n"
        "Используйте «❌ Отмена» и начните заново.",
        reply_markup=get_main_keyboard()
    )

# ----------------------------------------------------------------------
# ОБРАБОТЧИКИ СОСТОЯНИЙ (ОПРОС)
@dp.message_handler(state=Form.full_name)
async def process_full_name(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    if user_id not in user_data:
        await message.answer("❌ Ошибка: данные потеряны. Нажмите «📝 Заполнить анкету» заново.", reply_markup=get_main_keyboard())
        await state.finish()
        return
    user_data[user_id]['full_name'] = message.text.strip()
    await message.answer("Ваша <b>должность</b>:")
    await Form.next()

@dp.message_handler(state=Form.position)
async def process_position(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    if user_id not in user_data:
        await message.answer("❌ Ошибка: начните анкету заново.", reply_markup=get_main_keyboard())
        await state.finish()
        return
    user_data[user_id]['position'] = message.text.strip()
    await message.answer("Название <b>отдела / подразделения</b>:")
    await Form.next()

@dp.message_handler(state=Form.department)
async def process_department(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    if user_id not in user_data:
        await message.answer("❌ Ошибка: начните анкету заново.", reply_markup=get_main_keyboard())
        await state.finish()
        return
    user_data[user_id]['department'] = message.text.strip()
    await message.answer("📅 <b>Дата заполнения</b>:", reply_markup=get_date_keyboard())
    await Form.next()

@dp.callback_query_handler(lambda c: c.data.startswith('date_'), state=Form.date)
async def process_date_callback(callback: types.CallbackQuery, state: FSMContext):
    user_id = callback.from_user.id
    if user_id not in user_data:
        await callback.message.edit_text("❌ Ошибка: начните анкету заново.")
        await state.finish()
        await callback.answer()
        return
    if callback.data == "date_today":
        date_str = datetime.now().strftime("%d.%m.%Y")
        user_data[user_id]['date'] = date_str
        await callback.message.edit_text(f"📅 Дата выбрана: {date_str}")
        await bot.send_message(user_id, "Ваш <b>стаж в должности</b> (например, «2 года 5 месяцев»):")
        await Form.next()
    else:
        await callback.message.edit_text("Введите дату в формате <b>ДД.ММ.ГГГГ</b> (например, 25.12.2025):")
        await callback.answer()
        return
    await callback.answer()

@dp.message_handler(state=Form.date)
async def process_date_manual(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    if user_id not in user_data:
        await message.answer("❌ Ошибка: начните анкету заново.", reply_markup=get_main_keyboard())
        await state.finish()
        return
    user_data[user_id]['date'] = message.text.strip()
    await message.answer("Ваш <b>стаж в должности</b>:")
    await Form.next()

@dp.message_handler(state=Form.experience)
async def process_experience(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    if user_id not in user_data:
        await message.answer("❌ Ошибка: начните анкету заново.", reply_markup=get_main_keyboard())
        await state.finish()
        return
    user_data[user_id]['experience'] = message.text.strip()
    await message.answer("ФИО <b>непосредственного руководителя</b>:")
    await Form.next()

@dp.message_handler(state=Form.supervisor)
async def process_supervisor(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    if user_id not in user_data:
        await message.answer("❌ Ошибка: начните анкету заново.", reply_markup=get_main_keyboard())
        await state.finish()
        return
    user_data[user_id]['supervisor'] = message.text.strip()
    await message.answer("🎯 <b>Основная цель вашей работы</b> (1-2 предложения):")
    await Form.next()

@dp.message_handler(state=Form.mission_goal)
async def process_mission_goal(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    if user_id not in user_data:
        await message.answer("❌ Ошибка: начните анкету заново.", reply_markup=get_main_keyboard())
        await state.finish()
        return
    user_data[user_id]['mission_goal'] = message.text.strip()
    await message.answer("Как ваша работа <b>влияет на успех компании</b>?")
    await Form.next()

@dp.message_handler(state=Form.mission_impact)
async def process_mission_impact(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    if user_id not in user_data:
        await message.answer("❌ Ошибка: начните анкету заново.", reply_markup=get_main_keyboard())
        await state.finish()
        return
    user_data[user_id]['mission_impact'] = message.text.strip()
    await message.answer("Что было бы, <b>если бы вашей должности не существовало</b>?")
    await Form.next()

@dp.message_handler(state=Form.mission_absence)
async def process_mission_absence(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    if user_id not in user_data:
        await message.answer("❌ Ошибка: начните анкету заново.", reply_markup=get_main_keyboard())
        await state.finish()
        return
    user_data[user_id]['mission_absence'] = message.text.strip()
    await message.answer("📋 <b>Ключевые функции</b>\nОпишите 2-3 важнейшие функции. Для каждой: что входит, кому передаётся результат, стандарты.")
    await Form.next()

@dp.message_handler(state=Form.key_functions)
async def process_key_functions(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    if user_id not in user_data:
        await message.answer("❌ Ошибка: начните анкету заново.", reply_markup=get_main_keyboard())
        await state.finish()
        return
    user_data[user_id]['key_functions'] = message.text.strip()
    await message.answer("🧠 <b>Необходимые компетенции</b>\nПеречислите профессиональные, программные и мягкие навыки с уровнем владения (1-5) и примерами.")
    await Form.next()

@dp.message_handler(state=Form.competencies)
async def process_competencies(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    if user_id not in user_data:
        await message.answer("❌ Ошибка: начните анкету заново.", reply_markup=get_main_keyboard())
        await state.finish()
        return
    user_data[user_id]['competencies'] = message.text.strip()
    await message.answer("💰 <b>Прямая ценность вашей работы</b>\nПриведите измеримые показатели: экономия времени (часов/месяц), снижение ошибок (%), оптимизация затрат (руб.).")
    await Form.next()

@dp.message_handler(state=Form.direct_value)
async def process_direct_value(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    if user_id not in user_data:
        await message.answer("❌ Ошибка: начните анкету заново.", reply_markup=get_main_keyboard())
        await state.finish()
        return
    user_data[user_id]['direct_value'] = message.text.strip()
    await message.answer("🌟 <b>Косвенная ценность</b>\nЧто вы даёте компании, что нельзя измерить цифрами? (уникальные знания, поддержка процессов, связь между отделами)")
    await Form.next()

@dp.message_handler(state=Form.indirect_value)
async def process_indirect_value(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    if user_id not in user_data:
        await message.answer("❌ Ошибка: начните анкету заново.", reply_markup=get_main_keyboard())
        await state.finish()
        return
    user_data[user_id]['indirect_value'] = message.text.strip()
    await message.answer("⚠️ <b>Проблемы и сложности</b>\nЧто мешает работать эффективно? (технические, организационные, коммуникационные)")
    await Form.next()

@dp.message_handler(state=Form.problems)
async def process_problems(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    if user_id not in user_data:
        await message.answer("❌ Ошибка: начните анкету заново.", reply_markup=get_main_keyboard())
        await state.finish()
        return
    user_data[user_id]['problems'] = message.text.strip()
    await message.answer("💡 <b>Идеи по улучшению</b>\n1) Что можете улучшить сами?\n2) Что нужно улучшить с помощью компании?\n3) Инновационные идеи.")
    await Form.next()

@dp.message_handler(state=Form.improvements)
async def process_improvements(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    if user_id not in user_data:
        await message.answer("❌ Ошибка: начните анкету заново.", reply_markup=get_main_keyboard())
        await state.finish()
        return
    user_data[user_id]['improvements'] = message.text.strip()
    await message.answer("🎯 <b>Цели на ближайшие 6 месяцев</b>\nКакие цели ставите? Какие навыки хотите развить? Как видите свой рост?")
    await Form.next()

@dp.message_handler(state=Form.goals)
async def process_goals(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    if user_id not in user_data:
        await message.answer("❌ Ошибка: начните анкету заново.", reply_markup=get_main_keyboard())
        await state.finish()
        return
    user_data[user_id]['goals'] = message.text.strip()
    await message.answer("📢 <b>Обратная связь компании</b>\n1) Что работает хорошо (сохранить)?\n2) Что нужно изменить в компании/отделе?")
    await Form.next()

@dp.message_handler(state=Form.feedback)
async def process_feedback(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    if user_id not in user_data:
        await message.answer("❌ Ошибка: начните анкету заново.", reply_markup=get_main_keyboard())
        await state.finish()
        return
    user_data[user_id]['feedback'] = message.text.strip()
    await message.answer("⭐ <b>Оцените свою эффективность</b> по 10-балльной шкале (1 – очень низко, 10 – идеально).\nНажмите на кнопку:", reply_markup=get_rating_keyboard())
    await Form.next()

@dp.callback_query_handler(lambda c: c.data.startswith('rating_'), state=Form.self_rating)
async def process_rating_callback(callback: types.CallbackQuery, state: FSMContext):
    user_id = callback.from_user.id
    if user_id not in user_data:
        await callback.message.edit_text("❌ Ошибка: начните анкету заново.")
        await state.finish()
        await callback.answer()
        return
    rating = callback.data.split('_')[1]
    user_data[user_id]['self_rating'] = rating
    await callback.message.edit_text(f"⭐ Ваша оценка: {rating} из 10")
    await callback.answer()
    await bot.send_message(user_id, "💪 <b>Ваша главная сильная сторона</b> в работе:")
    await Form.next()

@dp.message_handler(state=Form.self_rating)
async def process_rating_text(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    if user_id not in user_data:
        await message.answer("❌ Ошибка: начните анкету заново.", reply_markup=get_main_keyboard())
        await state.finish()
        return
    user_data[user_id]['self_rating'] = message.text.strip()
    await message.answer("💪 <b>Ваша главная сильная сторона</b> в работе:")
    await Form.next()

@dp.message_handler(state=Form.self_strength)
async def process_self_strength(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    if user_id not in user_data:
        await message.answer("❌ Ошибка: начните анкету заново.", reply_markup=get_main_keyboard())
        await state.finish()
        return
    user_data[user_id]['self_strength'] = message.text.strip()
    await message.answer("📉 <b>Над чем нужно работать</b> (что стоит улучшить)?")
    await Form.next()

@dp.message_handler(state=Form.self_weakness)
async def process_self_weakness(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    if user_id not in user_data:
        await message.answer("❌ Ошибка: начните анкету заново.", reply_markup=get_main_keyboard())
        await state.finish()
        return
    user_data[user_id]['self_weakness'] = message.text.strip()
    await message.answer("🚀 <b>Одно изменение, которое больше всего повысит вашу эффективность</b>:")
    await Form.next()

@dp.message_handler(state=Form.self_one_change)
async def process_self_one_change(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    if user_id not in user_data:
        await message.answer("❌ Ошибка: начните анкету заново.", reply_markup=get_main_keyboard())
        await state.finish()
        return
    user_data[user_id]['self_one_change'] = message.text.strip()
    await message.answer("Спасибо! Формирую документ... Пожалуйста, подождите.")

    try:
        filename = generate_docx(user_data[user_id], user_id)
        with open(filename, 'rb') as doc_file:
            await message.answer_document(
                document=InputFile(doc_file, filename=filename),
                caption="✅ Ваша анкета готова!",
                reply_markup=get_main_keyboard()
            )
        os.remove(filename)
    except Exception as e:
        logger.exception("Ошибка при создании документа")
        await message.answer(f"❌ Ошибка: {e}. Попробуйте ещё раз, нажав «📝 Заполнить анкету».")

    # Очистка
    if user_id in user_data:
        del user_data[user_id]
    await state.finish()

# ----------------------------------------------------------------------
# ГЕНЕРАЦИЯ ДОКУМЕНТА .DOCX
def generate_docx(data: Dict[str, str], user_id: int) -> str:
    fio = data.get("full_name", "Сотрудник").replace(" ", "_")
    filename = f"Анкета_{fio}_{user_id}.docx"
    doc = Document()

    # Заголовок
    title = doc.add_heading("Анкета сотрудника", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fill_date = data.get("date", datetime.now().strftime("%d.%m.%Y"))
    doc.add_paragraph(f"Дата заполнения: {fill_date}")
    doc.add_paragraph()

    # 1. Основная информация
    doc.add_heading("1. Основная информация", level=1)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    table.cell(0,0).text = "ФИО сотрудника"
    table.cell(0,1).text = data.get("full_name", "")
    table.cell(1,0).text = "Должность"
    table.cell(1,1).text = data.get("position", "")
    table.cell(2,0).text = "Отдел/Подразделение"
    table.cell(2,1).text = data.get("department", "")
    table.cell(3,0).text = "Дата заполнения"
    table.cell(3,1).text = fill_date
    table.cell(4,0).text = "Стаж в должности"
    table.cell(4,1).text = data.get("experience", "")
    table.cell(5,0).text = "Непосредственный руководитель"
    table.cell(5,1).text = data.get("supervisor", "")

    # 2. Миссия
    doc.add_heading("2. Миссия моей должности", level=1)
    doc.add_paragraph(f"Основная цель работы:\n{data.get('mission_goal', '')}")
    doc.add_paragraph(f"Влияние на успех компании:\n{data.get('mission_impact', '')}")
    doc.add_paragraph(f"Если бы должности не существовало:\n{data.get('mission_absence', '')}")

    # 3. Ключевые функции
    doc.add_heading("3. Ключевые функции", level=1)
    doc.add_paragraph(data.get("key_functions", ""))

    # 4. Компетенции
    doc.add_heading("4. Необходимые компетенции", level=1)
    doc.add_paragraph(data.get("competencies", ""))

    # 5. Ценность
    doc.add_heading("5. Ценность моей работы для компании", level=1)
    doc.add_paragraph("Прямая ценность (измеримая):\n" + data.get("direct_value", ""))
    doc.add_paragraph("Косвенная ценность (неизмеримая):\n" + data.get("indirect_value", ""))

    # 6. Проблемы
    doc.add_heading("6. Проблемы и сложности", level=1)
    doc.add_paragraph(data.get("problems", ""))

    # 7. Идеи улучшений
    doc.add_heading("7. Идеи по улучшению", level=1)
    doc.add_paragraph(data.get("improvements", ""))

    # 8. Цели и развитие
    doc.add_heading("8. Мои цели и развитие", level=1)
    doc.add_paragraph(data.get("goals", ""))

    # 9. Обратная связь
    doc.add_heading("9. Обратная связь и предложения", level=1)
    doc.add_paragraph(data.get("feedback", ""))

    # 10. Самооценка
    doc.add_heading("10. Итоговая самооценка", level=1)
    doc.add_paragraph(f"Оценка эффективности (1-10):\n{data.get('self_rating', '')}")
    doc.add_paragraph(f"Главная сильная сторона:\n{data.get('self_strength', '')}")
    doc.add_paragraph(f"Над чем нужно работать:\n{data.get('self_weakness', '')}")
    doc.add_paragraph(f"Одно изменение для повышения эффективности:\n{data.get('self_one_change', '')}")

    doc.save(filename)
    return filename

# ----------------------------------------------------------------------
# ФОНОВЫЙ HTTP-СЕРВЕР ДЛЯ HEALTH CHECK (Render)
async def health_check(request):
    return web.Response(text="OK", status=200)

async def start_health_server():
    app = web.Application()
    app.router.add_get('/health', health_check)
    runner = web.AppRunner(app)
    await runner.setup()
    site = web.TCPSite(runner, '0.0.0.0', 8080)
    await site.start()
    logger.info("Health check server started on port 8080")

# ----------------------------------------------------------------------
# ЗАПУСК ПОЛЛИНГА + HEALTH-СЕРВЕР
async def on_startup(dp):
    asyncio.create_task(start_health_server())
    logger.info("Бот запущен в режиме polling")

async def on_shutdown(dp):
    await dp.storage.close()
    await dp.storage.wait_closed()

if __name__ == "__main__":
    executor.start_polling(dp, skip_updates=True, on_startup=on_startup, on_shutdown=on_shutdown)
